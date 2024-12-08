import sqlite3
from PyQt6.QtWidgets import QWidget, QLineEdit, QDateEdit, QComboBox, QPushButton, QFileDialog, QMessageBox
from datetime import datetime
from docx import Document
from random import randint

class Polic_form(QWidget):
    def __init__(self, login):
        self.item = None
        self.login = login
        super().__init__()
        self.setGeometry(0, 0, 200, 200)
        self.name = QLineEdit(self)
        self.calendar = QDateEdit(self)
        self.pol = QComboBox(self)
        self.form_button = QPushButton(self)
        self.setUI()

    def setUI(self):
        self.name.setPlaceholderText('Введите ФИО')
        self.calendar.setCalendarPopup(True)
        self.pol.addItem('муж')
        self.pol.addItem('жен')
        self.form_button.setText('Сформировать')
        self.form_button.clicked.connect(self.word)
        self.name.setGeometry(10, 10, 180, 30)
        self.calendar.setGeometry(10, 55, 180, 30)
        self.pol.setGeometry(10, 100, 180, 30)
        self.form_button.setGeometry(35, 140, 130, 40)

    def word(self):
        name = self.name.text()
        date = '-'.join(list(reversed(self.calendar.text().split('.'))))
        pol = self.pol.currentText()
        con = sqlite3.connect(r'../Users/lulun/PycharmProjects/PythonProject1/insurance.sqlite')
        kursor = con.cursor()

        kursor.execute(f'''UPDATE Client SET
                            full_name = ?,
                            gender = ?,
                            date_of_birth = ?
                            WHERE login = ?''', (name, pol, date, self.login))

        con.commit()
        con.close()

        code = []
        for i in range(16):
            code.append(str(randint(1, 9)))

        data = [['Персональный номер', ''.join(code)],
                ['ФИО', name],
                ['Пол', pol],
                ['Дата начала полиса', str(datetime.now())]
                ]

        doc = Document()#####

        doc.add_heading('ПОЛИС ОБЯЗАТЕЛЬНОГО МЕДИЦИНСКОГО СТРАХОВАНИЯ')

        table = doc.add_table(4, len(data[0]))
        table.style = 'Table Grid'
        for i in range(len(data)):
            for z in range(len(data[i])):
                cells = table.cell(i, z)
                cells.text = str(data[i][z])
        path, _ = QFileDialog.getSaveFileName(self, 'Save Excel', '.', 'Excel(*.docx)')
        if not path:
            self.item = QMessageBox.information(self, 'Внимание', 'Не указан файл для сохранения.')
            return
        doc.save(path)
